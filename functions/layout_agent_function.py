# coding:utf-8
"""
排版Agent功能组件
用于处理文档排版和Word文档生成
"""

import os
from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QPushButton, 
                             QLabel, QFileDialog, QTextEdit, QGroupBox, 
                             QProgressBar, QMessageBox, QFrame)
from PyQt6.QtCore import Qt, pyqtSignal
from qfluentwidgets import (PushButton, ComboBox, LineEdit, TextEdit, 
                           InfoBar, InfoBarPosition, PrimaryPushButton, 
                           ToggleButton, FluentIcon)
from .base_function import BaseFunction
# 直接复制layout_agent的核心功能到这个文件中
class LayoutAgent:
    """排版Agent类，负责文本结构分析和Word文档生成"""
    
    def __init__(self):
        pass
    
    def analyze_text_structure(self, text: str) -> dict:
        """使用简单文本分析"""
        return self.simple_text_analysis(text)
    
    def simple_text_analysis(self, text: str) -> dict:
        """简单的文本结构分析"""
        lines = text.split("\n")
        sections = []
        current_list = None
        current_table = None
        
        for line in lines:
            line = line.strip()
            if not line:
                # 处理段落分隔
                if current_list:
                    sections.append({
                        "type": "list",
                        "items": current_list
                    })
                    current_list = None
                if current_table:
                    sections.append({
                        "type": "table",
                        "rows": current_table
                    })
                    current_table = None
                continue
            
            # 检查是否为表格分隔线（Markdown表格格式）
            if line.startswith("|") and line.endswith("|") and all(c in "|- " for c in line):
                # 检查是否为表头分隔线
                if "---" in line:
                    # 跳过表头分隔线
                    continue
            
            # 检查是否为表格行
            if line.startswith("|") and line.endswith("|"):
                # 解析表格行
                cells = [cell.strip() for cell in line[1:-1].split("|")]
                if not current_table:
                    current_table = [cells]
                else:
                    current_table.append(cells)
                continue
            
            # 检查是否为标题（以#开头）
            if line.startswith("#"):
                # 结束当前列表或表格
                if current_list:
                    sections.append({
                        "type": "list",
                        "items": current_list
                    })
                    current_list = None
                if current_table:
                    sections.append({
                        "type": "table",
                        "rows": current_table
                    })
                    current_table = None
                
                # 计算标题级别
                level = len(line.split()[0])
                content = line[level:].strip()
                sections.append({
                    "type": "heading",
                    "level": level,
                    "content": content
                })
            # 检查是否为列表项（以-、*、+或数字开头）
            elif line.startswith(("-", "*", "+")) or any(line.startswith(f"{i}.") for i in range(1, 10)):
                # 结束当前表格
                if current_table:
                    sections.append({
                        "type": "table",
                        "rows": current_table
                    })
                    current_table = None
                
                # 获取列表内容
                if line.startswith(("-", "*", "+")):
                    content = line[1:].strip()
                else:
                    # 数字列表
                    content = line.split(".", 1)[1].strip()
                
                # 开始新列表或添加到现有列表
                if not current_list:
                    current_list = [content]
                else:
                    current_list.append(content)
            else:
                # 结束当前列表或表格
                if current_list:
                    sections.append({
                        "type": "list",
                        "items": current_list
                    })
                    current_list = None
                if current_table:
                    sections.append({
                        "type": "table",
                        "rows": current_table
                    })
                    current_table = None
                
                # 普通段落
                sections.append({
                    "type": "paragraph",
                    "content": line
                })
        
        # 处理最后一个列表或表格
        if current_list:
            sections.append({
                "type": "list",
                "items": current_list
            })
        if current_table:
            sections.append({
                "type": "table",
                "rows": current_table
            })
        
        # 尝试从第一段中提取标题
        title = ""
        if sections and sections[0]["type"] == "heading" and sections[0]["level"] == 1:
            # 如果第一段是一级标题，使用它作为文档标题
            title = sections[0]["content"]
        elif sections and sections[0]["type"] == "paragraph":
            # 否则尝试从第一段提取可能的标题
            first_para = sections[0]["content"]
            if len(first_para) < 50 and any(char.isupper() for char in first_para):
                title = first_para
        
        return {
            "title": title,
            "sections": sections
        }
    
    def generate_word_document(self, structure: dict, output_path: str) -> tuple:
        """根据文本结构生成Word文档"""
        try:
            # 创建新文档
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 在文档开头添加目录
            doc.add_heading("目录", level=1)
            # 添加目录字段
            from docx.oxml.shared import OxmlElement, qn
            
            # 创建目录段落
            toc_paragraph = doc.add_paragraph()
            
            # 创建目录字段
            run = toc_paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = 'TOC o "1-3" h z u'
            run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            run._r.append(fldChar2)
            
            fldChar3 = OxmlElement('w:t')
            fldChar3.text = '更新目录'
            run._r.append(fldChar3)
            
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar4)
            
            # 添加分页符
            doc.add_page_break()
            
            # 添加标题 - 只在标题非空时添加
            title = structure.get("title", "")
            if title:  # 只在标题非空时添加
                doc.add_heading(title, level=0)
            
            # 添加章节
            sections = structure.get("sections", [])
            for section in sections:
                if section["type"] == "heading":
                    # 添加标题
                    level = section.get("level", 1)
                    # 确保级别在1-9之间
                    level = max(1, min(9, level))
                    doc.add_heading(section["content"], level=level)
                elif section["type"] == "paragraph":
                    # 添加段落
                    para = doc.add_paragraph(section["content"])
                    # 设置段落格式
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(12)
                    para_format.line_spacing = 1.5
                elif section["type"] == "list":
                    # 添加列表
                    items = section.get("items", [])
                    for item in items:
                        doc.add_paragraph(item, style="List Bullet")
                elif section["type"] == "table":
                    # 添加表格
                    rows = section.get("rows", [])
                    if not rows:
                        continue
                    
                    # 获取表格行列数
                    num_rows = len(rows)
                    num_cols = len(rows[0])
                    
                    # 创建表格
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    
                    # 设置表格居中对齐
                    try:
                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
                    
                    # 自动调整表格宽高
                    try:
                        # 计算可用宽度
                        page_width = doc.sections[0].page_width
                        margins = doc.sections[0].left_margin + doc.sections[0].right_margin
                        available_width = page_width - margins
                        table_width = available_width * 0.9
                        
                        # 计算每列宽度（均匀分配）
                        col_width = table_width / num_cols
                        
                        # 设置统一的行高
                        row_height = Pt(20)  # 可以根据需要调整
                        
                        # 添加表格内容并美化
                        for i, row in enumerate(rows):
                            table_row = table.rows[i]
                            table_row.height = row_height
                            table_row.height_rule = True
                            
                            for j, cell_text in enumerate(row):
                                cell = table_row.cells[j]
                                cell.text = cell_text
                                
                                # 设置单元格宽度（均匀分配）
                                cell.width = col_width
                                
                                # 设置单元格格式
                                paragraph = cell.paragraphs[0]
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.space_after = Pt(0)
                                paragraph_format.line_spacing = 1.0
                                
                                # 统一字体样式
                                for run in paragraph.runs:
                                    run.font.name = '微软雅黑'  # 设置统一字体
                                    run.font.size = Pt(10)  # 统一字体大小
                                    
                                    # 表头仍然保持粗体
                                    if i == 0:
                                        run.font.bold = True
                                    else:
                                        run.font.bold = False
                                
                                # 设置单元格对齐方式 - 居中对齐
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Error formatting table: {str(e)}")
                    
                    # 为表格添加边框
                    try:
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        
                        # 获取表格XML元素
                        tbl = table._tbl
                        
                        # 创建边框属性
                        border_props = [
                            ('top', 'single', '0.5pt', 'auto'),
                            ('bottom', 'single', '0.5pt', 'auto'),
                            ('left', 'single', '0.5pt', 'auto'),
                            ('right', 'single', '0.5pt', 'auto')
                        ]
                        
                        # 为每个单元格添加边框
                        for row in tbl.iterchildren():
                            for cell in row.iterchildren():
                                tc = cell
                                tcPr = tc.get_or_add_tcPr()
                                
                                for border_name, border_type, border_size, border_color in border_props:
                                    border = OxmlElement(f'w:{border_name}')
                                    border.set(qn('w:val'), border_type)
                                    border.set(qn('w:sz'), border_size)
                                    border.set(qn('w:space'), '0')
                                    border.set(qn('w:color'), border_color)
                                    tcPr.append(border)
                    except Exception:
                        pass
            
            # 保存文档
            doc.save(output_path)
            return True, ""
        except Exception as e:
            error_msg = f"Error generating Word document: {str(e)}"
            print(error_msg)
            return False, error_msg
    
    def process_text(self, text: str, output_path: str) -> tuple:
        """处理文本，生成排版后的Word文档"""
        try:
            # 直接排版文本，不进行模型分析，完全保留原文内容，包括表格
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建新文档
            doc = Document()
            
            # 使用simple_text_analysis来检测文本结构，特别是表格
            structure = self.simple_text_analysis(text)
            sections = structure.get("sections", [])
            
            # 遍历所有检测到的内容块
            for section in sections:
                if section["type"] == "heading":
                    # 添加标题
                    level = section.get("level", 1)
                    level = max(1, min(9, level))
                    doc.add_heading(section["content"], level=level)
                elif section["type"] == "paragraph":
                    # 添加段落
                    para = doc.add_paragraph(section["content"])
                    # 设置段落格式
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(12)
                    para_format.line_spacing = 1.5
                elif section["type"] == "list":
                    # 添加列表
                    items = section.get("items", [])
                    for item in items:
                        doc.add_paragraph(item, style="List Bullet")
                elif section["type"] == "table":
                    # 添加表格
                    rows = section.get("rows", [])
                    if not rows:
                        continue
                    
                    # 获取表格行列数
                    num_rows = len(rows)
                    num_cols = len(rows[0])
                    
                    # 创建表格
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    
                    # 设置表格居中对齐
                    try:
                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
                    
                    # 自动调整表格宽高
                    try:
                        # 计算可用宽度
                        page_width = doc.sections[0].page_width
                        margins = doc.sections[0].left_margin + doc.sections[0].right_margin
                        available_width = page_width - margins
                        table_width = available_width * 0.9
                        
                        # 计算每列宽度（均匀分配）
                        col_width = table_width / num_cols
                        
                        # 设置统一的行高
                        row_height = Pt(20)  # 可以根据需要调整
                        
                        # 添加表格内容并美化
                        for i, row in enumerate(rows):
                            table_row = table.rows[i]
                            table_row.height = row_height
                            table_row.height_rule = True
                            
                            for j, cell_text in enumerate(row):
                                cell = table_row.cells[j]
                                cell.text = cell_text
                                
                                # 设置单元格宽度（均匀分配）
                                cell.width = col_width
                                
                                # 设置单元格格式
                                paragraph = cell.paragraphs[0]
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.space_after = Pt(0)
                                paragraph_format.line_spacing = 1.0
                                
                                # 统一字体样式
                                for run in paragraph.runs:
                                    run.font.name = '微软雅黑'  # 设置统一字体
                                    run.font.size = Pt(10)  # 统一字体大小
                                    
                                    # 表头仍然保持粗体
                                    if i == 0:
                                        run.font.bold = True
                                    else:
                                        run.font.bold = False
                                
                                # 设置单元格对齐方式 - 居中对齐
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Error formatting table: {str(e)}")
                    
                    # 为表格添加边框
                    try:
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        
                        # 获取表格XML元素
                        tbl = table._tbl
                        
                        # 创建边框属性
                        border_props = [
                            ('top', 'single', '0.5pt', 'auto'),
                            ('bottom', 'single', '0.5pt', 'auto'),
                            ('left', 'single', '0.5pt', 'auto'),
                            ('right', 'single', '0.5pt', 'auto')
                        ]
                        
                        # 为每个单元格添加边框
                        for row in tbl.iterchildren():
                            for cell in row.iterchildren():
                                tc = cell
                                tcPr = tc.get_or_add_tcPr()
                                
                                for border_name, border_type, border_size, border_color in border_props:
                                    border = OxmlElement(f'w:{border_name}')
                                    border.set(qn('w:val'), border_type)
                                    border.set(qn('w:sz'), border_size)
                                    border.set(qn('w:space'), '0')
                                    border.set(qn('w:color'), border_color)
                                    tcPr.append(border)
                    except Exception:
                        pass
            
            # 保存文档
            doc.save(output_path)
            return True, "文本排版完成，完全保留了原文内容，包括表格"
        except Exception as e:
            error_msg = f"Error processing text: {str(e)}"
            print(error_msg)
            return False, error_msg
    
    def process_file(self, input_path: str, output_path: str) -> tuple:
        """处理文件，生成排版后的Word文档"""
        try:
            # 读取文件内容
            import os
            file_ext = os.path.splitext(input_path)[1].lower()
            
            if file_ext in ['.docx']:
                # 直接操作原始文档，只进行排版和美化
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # 打开原文档
                doc = Document(input_path)
                
                # 1. 美化所有段落格式，但特殊处理表格前的段落（包括两行标题的情况）
                for i, para in enumerate(doc.paragraphs):
                    # 设置段落格式
                    para_format = para.paragraph_format
                    para_format.line_spacing = 1.5
                    
                    # 检查是否为表格前的段落
                    is_table_title = False
                    is_table_related = False
                    
                    # 检查段落是否与表格相关
                    for table in doc.tables:
                        # 获取表格的位置
                        table_element = table._element
                        para_element = para._element
                        
                        # 情况1：段落后面直接跟着表格
                        if para_element.getnext() == table_element:
                            is_table_title = True
                            is_table_related = True
                            break
                        
                        # 情况2：段落后面跟着另一个段落，而那个段落后面跟着表格（两行标题的情况）
                        if i < len(doc.paragraphs) - 1:
                            next_para_element = doc.paragraphs[i + 1]._element
                            if next_para_element.getnext() == table_element:
                                is_table_related = True
                                break
                    
                    # 设置间距
                    if is_table_title:
                        # 表格直接前的段落（通常是单位行），与表格完全贴合
                        para_format.space_after = Pt(0)
                    elif is_table_related:
                        # 表格前的前一个段落（通常是主标题），减小与单位行的间距
                        para_format.space_after = Pt(1)
                    else:
                        # 普通段落间距调低
                        para_format.space_after = Pt(8)
                
                # 2. 美化所有表格
                for table in doc.tables:
                    # 设置表格居中对齐
                    try:
                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
                    
                    # 自动调整表格宽高
                    try:
                        # 计算可用宽度
                        page_width = doc.sections[0].page_width
                        margins = doc.sections[0].left_margin + doc.sections[0].right_margin
                        available_width = page_width - margins
                        table_width = available_width * 0.9
                        
                        # 计算每列宽度（均匀分配）
                        num_cols = len(table.columns)
                        col_width = table_width / num_cols
                        
                        # 设置统一的行高
                        row_height = Pt(20)  # 可以根据需要调整
                        
                        # 美化表格内容
                        for i, row in enumerate(table.rows):
                            # 设置统一行高
                            row.height = row_height
                            row.height_rule = True
                            
                            for cell in row.cells:
                                # 设置单元格宽度（均匀分配）
                                cell.width = col_width
                                
                                # 设置单元格内段落格式
                                for para in cell.paragraphs:
                                    para_format = para.paragraph_format
                                    para_format.space_after = Pt(0)  # 单元格内段落间距为0
                                    para_format.line_spacing = 1.2  # 略微减小行高，更加紧凑
                                    
                                    # 统一字体样式
                                    for run in para.runs:
                                        run.font.name = '微软雅黑'  # 设置统一字体
                                        run.font.size = Pt(10)  # 统一字体大小
                                        
                                        # 表头仍然保持粗体
                                        if i == 0:
                                            run.font.bold = True
                                        else:
                                            run.font.bold = False
                                    
                                    # 设置单元格内容居中
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Error formatting table: {str(e)}")
                    
                    # 为表格添加边框
                    try:
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        
                        # 获取表格XML元素
                        tbl = table._tbl
                        
                        # 创建边框属性
                        border_props = [
                            ('top', 'single', '0.5pt', 'auto'),
                            ('bottom', 'single', '0.5pt', 'auto'),
                            ('left', 'single', '0.5pt', 'auto'),
                            ('right', 'single', '0.5pt', 'auto')
                        ]
                        
                        # 为每个单元格添加边框
                        for row in tbl.iterchildren():
                            for cell in row.iterchildren():
                                tc = cell
                                tcPr = tc.get_or_add_tcPr()
                                
                                for border_name, border_type, border_size, border_color in border_props:
                                    border = OxmlElement(f'w:{border_name}')
                                    border.set(qn('w:val'), border_type)
                                    border.set(qn('w:sz'), border_size)
                                    border.set(qn('w:space'), '0')
                                    border.set(qn('w:color'), border_color)
                                    tcPr.append(border)
                    except Exception:
                        pass
                
                # 保存美化后的文档
                doc.save(output_path)
                return True, "文档排版和表格美化完成，保留了所有原文内容"
            elif file_ext in ['.txt', '.md']:
                # 对于文本文件，创建文档并完全保留原文内容，包括表格，仅添加排版格式
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # 创建新文档
                doc = Document()
                
                # 读取文本内容
                with open(input_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                # 使用simple_text_analysis来检测文本结构，特别是表格
                structure = self.simple_text_analysis(text)
                sections = structure.get("sections", [])
                
                # 遍历所有检测到的内容块
                for section in sections:
                    if section["type"] == "heading":
                        # 添加标题
                        level = section.get("level", 1)
                        level = max(1, min(9, level))
                        doc.add_heading(section["content"], level=level)
                    elif section["type"] == "paragraph":
                        # 添加段落
                        para = doc.add_paragraph(section["content"])
                        # 设置段落格式
                        para_format = para.paragraph_format
                        para_format.space_after = Pt(12)
                        para_format.line_spacing = 1.5
                    elif section["type"] == "list":
                        # 添加列表
                        items = section.get("items", [])
                        for item in items:
                            doc.add_paragraph(item, style="List Bullet")
                    elif section["type"] == "table":
                        # 添加表格
                        rows = section.get("rows", [])
                        if not rows:
                            continue
                        
                        # 获取表格行列数
                        num_rows = len(rows)
                        num_cols = len(rows[0])
                        
                        # 创建表格
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        
                        # 设置表格居中对齐
                        try:
                            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception:
                            pass
                        
                        # 自动调整表格宽高
                        try:
                            # 计算可用宽度
                            page_width = doc.sections[0].page_width
                            margins = doc.sections[0].left_margin + doc.sections[0].right_margin
                            available_width = page_width - margins
                            table_width = available_width * 0.9
                            
                            # 计算每列宽度（均匀分配）
                            col_width = table_width / num_cols
                            
                            # 设置统一的行高
                            row_height = Pt(20)  # 可以根据需要调整
                            
                            # 添加表格内容并美化
                            for i, row in enumerate(rows):
                                table_row = table.rows[i]
                                table_row.height = row_height
                                table_row.height_rule = True
                                
                                for j, cell_text in enumerate(row):
                                    cell = table_row.cells[j]
                                    cell.text = cell_text
                                    
                                    # 设置单元格宽度（均匀分配）
                                    cell.width = col_width
                                    
                                    # 设置单元格格式
                                    paragraph = cell.paragraphs[0]
                                    paragraph_format = paragraph.paragraph_format
                                    paragraph_format.space_after = Pt(0)
                                    paragraph_format.line_spacing = 1.0
                                    
                                    # 统一字体样式
                                    for run in paragraph.runs:
                                        run.font.name = '微软雅黑'  # 设置统一字体
                                        run.font.size = Pt(10)  # 统一字体大小
                                        
                                        # 表头仍然保持粗体
                                        if i == 0:
                                            run.font.bold = True
                                        else:
                                            run.font.bold = False
                                    
                                    # 设置单元格对齐方式 - 居中对齐
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            print(f"Error formatting table: {str(e)}")
                
                # 保存文档
                doc.save(output_path)
                return True, "文本排版完成，包括表格，完全保留了原文内容"
            else:
                # 不支持的文件格式
                error_msg = f"Unsupported file format: {file_ext}"
                print(error_msg)
                return False, error_msg
        except Exception as e:
            error_msg = f"Error processing file: {str(e)}"
            print(error_msg)
            return False, error_msg

# 创建全局排版Agent实例
layout_agent = LayoutAgent()

class LayoutAgentFunction(BaseFunction):
    """文档排版功能组件"""
    
    def __init__(self, parent=None):
        description = (
            "📢 <b>功能说明：</b><br>" 
            "文档排版工具可以帮助您处理文档排版，生成美观的Word文档。<br>" 
            "支持文本直接输入和文件导入两种方式。"
        )
        super().__init__("文档排版", description, parent)
        self.init_ui()
    
    def init_ui(self):
        """初始化UI界面"""
        # 创建文件选择区域
        self.file_group = QGroupBox("文件处理")
        self.file_layout = QVBoxLayout(self.file_group)
        self.file_layout.setSpacing(10)
        
        # 输入文件选择
        self.input_file_layout = QHBoxLayout()
        self.input_file_edit = LineEdit()
        self.input_file_edit.setPlaceholderText("选择输入文件")
        self.browse_input_btn = PushButton("浏览", icon=FluentIcon.FOLDER)
        self.browse_input_btn.clicked.connect(self.browse_input_file)
        
        self.input_file_layout.addWidget(self.input_file_edit)
        self.input_file_layout.addWidget(self.browse_input_btn)
        
        # 输出文件设置
        self.output_file_layout = QHBoxLayout()
        self.output_file_edit = LineEdit()
        self.output_file_edit.setPlaceholderText("设置输出文件路径")
        self.browse_output_btn = PushButton("设置", icon=FluentIcon.SAVE)
        self.browse_output_btn.clicked.connect(self.browse_output_file)
        
        self.output_file_layout.addWidget(self.output_file_edit)
        self.output_file_layout.addWidget(self.browse_output_btn)
        
        # 添加到文件组
        self.file_layout.addLayout(self.input_file_layout)
        self.file_layout.addLayout(self.output_file_layout)
        
        # 创建文本输入区域
        self.text_group = QGroupBox("直接文本输入")
        self.text_layout = QVBoxLayout(self.text_group)
        
        self.text_edit = TextEdit()
        self.text_edit.setPlaceholderText("您可以直接输入文本内容，支持Markdown格式...")
        self.text_layout.addWidget(self.text_edit)
        
        # 创建操作按钮
        self.process_file_btn = PrimaryPushButton("处理文件")
        self.process_file_btn.clicked.connect(self.process_file)
        
        self.process_text_btn = PrimaryPushButton("处理文本")
        self.process_text_btn.clicked.connect(self.process_text)
        
        self.clear_btn = PushButton("清空")
        self.clear_btn.clicked.connect(self.clear_text)
        
        self.addExecuteButton("处理文件", self.process_file)
        self.buttonLayout.addWidget(self.process_text_btn)
        self.buttonLayout.addWidget(self.clear_btn)
        
        # 进度条容器
        self.progress_container = QWidget()
        self.progress_container.setFixedHeight(50)
        self.progress_layout = QVBoxLayout(self.progress_container)
        self.progress_layout.setContentsMargins(0, 0, 0, 0)
        self.progress_layout.setSpacing(5)
        
        # 进度文本
        self.progress_text = QLabel("准备开始排版...", self)
        self.progress_text.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # 进度条
        self.progress_bar = QFrame(self)
        self.progress_bar.setFixedHeight(4)
        self.progress_bar.setStyleSheet("""QFrame {
                background-color: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                    stop:0 #e0e0e0, stop:1 #e0e0e0);
                border-radius: 2px;
            }""")
        
        # 将进度文本和进度条添加到容器
        self.progress_layout.addWidget(self.progress_text)
        self.progress_layout.addWidget(self.progress_bar)
        
        # 设置容器初始不可见
        self.progress_container.setVisible(False)
        
        # 添加所有组件到内容布局
        self.contentLayout.addWidget(self.file_group)
        self.contentLayout.addWidget(self.text_group)
        self.contentLayout.addWidget(self.progress_container)
        
        # 应用初始主题
        self._onThemeChanged()
    
    def browse_input_file(self):
        """浏览输入文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择输入文件", "/", "支持的文件 (*.txt *.md *.docx)"
        )
        if file_path:
            self.input_file_edit.setText(file_path)
            # 自动生成输出文件路径
            output_path = os.path.splitext(file_path)[0] + "_排版.docx"
            self.output_file_edit.setText(output_path)
    
    def browse_output_file(self):
        """浏览输出文件"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存输出文件", "/", "Word文档 (*.docx)"
        )
        if file_path:
            # 确保文件扩展名为.docx
            if not file_path.endswith(".docx"):
                file_path += ".docx"
            self.output_file_edit.setText(file_path)
    
    def process_file(self):
        """处理文件"""
        input_path = self.input_file_edit.text().strip()
        output_path = self.output_file_edit.text().strip()
        
        if not input_path:
            QMessageBox.warning(self, "警告", "请选择输入文件！")
            return
        
        if not output_path:
            QMessageBox.warning(self, "警告", "请设置输出文件路径！")
            return
        
        # 显示进度条
        self.progress_container.setVisible(True)
        self.updateProgress(0, "正在准备处理文件...")
        
        try:
            # 调用排版Agent处理文件
            self.updateProgress(20, "正在读取文件内容...")
            success, message = layout_agent.process_file(input_path, output_path)
            
            if success:
                self.updateProgress(100, "文件处理完成！")
                self.showSuccess(f"文件处理完成！\n{message}")
            else:
                self.updateProgress(100, "文件处理失败！")
                self.showError(f"文件处理失败！\n{message}")
        except Exception as e:
            self.updateProgress(100, "处理文件时发生错误！")
            self.showError(f"处理文件时发生错误：{str(e)}")
        finally:
            self.progress_container.setVisible(False)
    
    def process_text(self):
        """处理文本"""
        text = self.text_edit.toPlainText().strip()
        output_path = self.output_file_edit.text().strip()
        
        if not text:
            QMessageBox.warning(self, "警告", "请输入文本内容！")
            return
        
        if not output_path:
            # 生成默认输出路径
            output_path = os.path.join(os.getcwd(), "排版结果.docx")
            self.output_file_edit.setText(output_path)
        
        # 显示进度条
        self.progress_container.setVisible(True)
        self.updateProgress(0, "正在准备处理文本...")
        
        try:
            # 调用排版Agent处理文本
            self.updateProgress(20, "正在处理文本内容...")
            success, message = layout_agent.process_text(text, output_path)
            
            if success:
                self.updateProgress(100, "文本处理完成！")
                self.showSuccess(f"文本处理完成！\n{message}")
            else:
                self.updateProgress(100, "文本处理失败！")
                self.showError(f"文本处理失败！\n{message}")
        except Exception as e:
            self.updateProgress(100, "处理文本时发生错误！")
            self.showError(f"处理文本时发生错误：{str(e)}")
        finally:
            self.progress_container.setVisible(False)
    
    def clear_text(self):
        """清空文本输入"""
        self.text_edit.clear()
    
    def execute(self):
        """执行功能（BaseFunction要求实现的方法）"""
        # 检查当前输入状态，自动选择合适的处理方法
        if self.input_file_edit.text().strip():
            # 如果有输入文件，调用文件处理方法
            self.process_file()
        elif self.text_edit.toPlainText().strip():
            # 如果有文本输入，调用文本处理方法
            self.process_text()
        else:
            # 没有输入，显示警告
            QMessageBox.warning(self, "警告", "请选择输入文件或输入文本内容！")